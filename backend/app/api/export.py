from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.database import get_db
from app.models import RPS
from app.schemas import ExportRequest
from sqlalchemy import select
import json
import os
from app.core.config import settings
from jinja2 import Template

router = APIRouter(prefix="/export", tags=["Export"])


RPS_HTML_TEMPLATE = Template("""
<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <title>RPS - {{ data.identitas.nama_mata_kuliah if data.identitas else '' }}</title>
  <style>
    @page {
      size: A4;
      margin: 1.5cm;
    }
    body { font-family: 'Times New Roman', serif; font-size: 10pt; color: #000; line-height: 1.3; }
    h1 { text-align: center; font-size: 13pt; margin-bottom: 5px; font-weight: bold; }
    h2 { font-size: 11pt; margin-top: 15px; margin-bottom: 5px; font-weight: bold; border-bottom: 1px solid #000; padding-bottom: 2px; }
    h3 { font-size: 10pt; margin-top: 8px; margin-bottom: 3px; font-weight: bold; }
    table { width: 100%; border-collapse: collapse; margin: 10px 0; table-layout: fixed; }
    th, td { border: 1px solid #000; padding: 5px; text-align: left; font-size: 9pt; word-wrap: break-word; vertical-align: top; }
    th { background: #f2f2f2; font-weight: bold; text-align: center; }
    .header-table td { padding: 8px; vertical-align: middle; }
    .text-center { text-align: center; }
    ul, ol { margin: 5px 0; padding-left: 18px; }
    li { font-size: 9pt; margin-bottom: 2px; }
  </style>
</head>
<body>
  
  <!-- Kop Sekolah / Perguruan Tinggi -->
  <table class="header-table" style="width: 100%; margin-bottom: 15px;">
    <tr>
      <td rowspan="2" style="width: 15%; text-align: center; border: 1px solid #000;">
        {% if brand_logo %}
          <img src="{{ brand_logo }}" style="max-height: 60px; max-width: 100%;" />
        {% else %}
          <div style="font-weight: bold; font-size: 8pt; color: #777;">LOGO</div>
        {% endif %}
      </td>
      <td style="width: 85%; text-align: center; border: 1px solid #000; font-weight: bold; font-size: 13pt; letter-spacing: 0.5px;">
        {{ brand_name }}
      </td>
    </tr>
    <tr>
      <td style="text-align: center; border: 1px solid #000; font-weight: bold; font-size: 11pt;">
        JURUSAN / PROGRAM STUDI {{ prodi_name | upper }}
      </td>
    </tr>
  </table>

  <!-- Judul dan Identitas RPS -->
  <table style="width: 100%; margin-bottom: 15px;">
    <tr>
      <td colspan="4" style="text-align: center; font-weight: bold; background-color: #f2f2f2; font-size: 11pt; padding: 6px;">
        RENCANA PEMBELAJARAN SEMESTER (RPS)
      </td>
    </tr>
    <tr>
      <td style="font-weight: bold; width: 25%;">Nama Mata Kuliah</td>
      <td style="width: 25%;">{{ data.identitas.nama_mata_kuliah if data.identitas else '' }}</td>
      <td style="font-weight: bold; width: 20%;">Kode MK</td>
      <td style="width: 30%;">{{ data.identitas.kode_mata_kuliah if data.identitas else '' }}</td>
    </tr>
    <tr>
      <td style="font-weight: bold;">Bobot (SKS)</td>
      <td>{{ data.identitas.sks if data.identitas else '' }} SKS</td>
      <td style="font-weight: bold;">Semester</td>
      <td>{{ data.identitas.semester if data.identitas else '' }}</td>
    </tr>
    <tr>
      <td style="font-weight: bold;">Tanggal Penyusunan</td>
      <td>{{ data.identitas.tanggal_penyusunan if data.identitas else '-' }}</td>
      <td style="font-weight: bold;">No Dokumen</td>
      <td>{{ data.identitas.no_dokumen if data.identitas else '-' }}</td>
    </tr>
    <tr>
      <td colspan="4" style="font-weight: bold; background-color: #f9f9f9; text-align: center; padding: 4px;">
        OTORISASI
      </td>
    </tr>
    <tr>
      <td colspan="2" style="font-weight: bold; text-align: center; padding: 4px;">Koordinator Pengembang RPS</td>
      <td style="font-weight: bold; text-align: center; padding: 4px;">Koordinator Rumpun MK</td>
      <td style="font-weight: bold; text-align: center; padding: 4px;">Ka PRODI</td>
    </tr>
    <tr style="height: 50px;">
      <td colspan="2" style="text-align: center; vertical-align: bottom; padding-bottom: 5px;">
        Tanda tangan<br/><strong>{{ data.identitas.koordinator_pengembang_rps if data.identitas else '-' }}</strong>
      </td>
      <td style="text-align: center; vertical-align: bottom; padding-bottom: 5px;">
        Tanda tangan<br/><strong>{{ data.identitas.koordinator_rmk if data.identitas else '-' }}</strong>
      </td>
      <td style="text-align: center; vertical-align: bottom; padding-bottom: 5px;">
        Tanda tangan<br/><strong>{{ data.identitas.ka_prodi if data.identitas else '-' }}</strong>
      </td>
    </tr>
    <tr>
      <td style="font-weight: bold; width: 25%;">Dosen Pengampu</td>
      <td colspan="3">
        {% if data.dosen_pengampu is iterable and data.dosen_pengampu is not string %}
          {% for d in data.dosen_pengampu %}
            {{ d.nama if d.nama else d }}{% if not loop.last %}, {% endif %}
          {% endfor %}
        {% else %}
          {{ data.dosen_pengampu or '-' }}
        {% endif %}
      </td>
    </tr>
    <tr>
      <td style="font-weight: bold;">Mata Kuliah Prasyarat</td>
      <td colspan="3">{{ data.prasyarat or '-' }}</td>
    </tr>
  </table>

  <h2>Capaian Pembelajaran (CP)</h2>
  
  <h3>1. CPL-PRODI (Capaian Pembelajaran Lulusan Program Studi) Yang Dibebankan Pada Mata Kuliah</h3>
  <table>
    <tr>
      <th width="15%">Kode CPL</th>
      <th width="85%">Deskripsi Capaian Pembelajaran</th>
    </tr>
    {% for cpl in course_cpls %}
    <tr>
      <td style="font-weight: bold; text-align: center;">{{ cpl.kode }}</td>
      <td>{{ cpl.deskripsi }}</td>
    </tr>
    {% else %}
    <tr>
      <td colspan="2" class="text-center" style="color: #777; font-style: italic;">Belum ada pemetaan CPL prodi yang dispesifikasikan.</td>
    </tr>
    {% endfor %}
  </table>

  <h3>2. CPMK (Capaian Pembelajaran Mata Kuliah)</h3>
  <table>
    <tr>
      <th width="15%">Kode CPMK</th>
      <th width="65%">Deskripsi CPMK</th>
      <th width="20%">CPL Yang Didukung</th>
    </tr>
    {% for c in data.cpmk %}
    {% if c is mapping %}
    <tr>
      <td style="font-weight: bold; text-align: center;">{{ c.kode }}</td>
      <td>{{ c.deskripsi }}</td>
      <td>
        {% if c.cpl_prodi is string %}
          {{ c.cpl_prodi }}
        {% elif c.cpl_prodi is iterable %}
          {{ c.cpl_prodi | join(', ') }}
        {% else %}
          {{ c.cpl_prodi }}
        {% endif %}
      </td>
    </tr>
    {% endif %}
    {% endfor %}
  </table>

  <h2>Deskripsi Singkat Mata Kuliah</h2>
  <p style="text-align: justify; margin: 5px 0;">{{ data.deskripsi_mata_kuliah }}</p>

  <h2>Bahan Kajian / Pokok Bahasan</h2>
  {% if data.bahan_kajian is iterable and data.bahan_kajian is not string %}
    <ul>
      {% for bk in data.bahan_kajian %}
        <li>{{ bk }}</li>
      {% endfor %}
    </ul>
  {% else %}
    <p>{{ data.bahan_kajian or '-' }}</p>
  {% endif %}

  <h2>Media Pembelajaran</h2>
  <table>
    <tr>
      <th width="50%">Perangkat Lunak (Software)</th>
      <th width="50%">Perangkat Keras (Hardware)</th>
    </tr>
    <tr>
      <td>
        {% if data.media_pembelajaran is mapping %}
          {% if data.media_pembelajaran.perangkat_lunak is iterable and data.media_pembelajaran.perangkat_lunak is not string %}
            <ul>
              {% for s in data.media_pembelajaran.perangkat_lunak %}
                <li>{{ s }}</li>
              {% endfor %}
            </ul>
          {% else %}
            {{ data.media_pembelajaran.perangkat_lunak or '-' }}
          {% endif %}
        {% else %}
          {% if data.media_pembelajaran is iterable and data.media_pembelajaran is not string %}
            <ul>
              {% for m in data.media_pembelajaran %}
                <li>{{ m }}</li>
              {% endfor %}
            </ul>
          {% else %}
            {{ data.media_pembelajaran or '-' }}
          {% endif %}
        {% endif %}
      </td>
      <td>
        {% if data.media_pembelajaran is mapping %}
          {% if data.media_pembelajaran.perangkat_keras is iterable and data.media_pembelajaran.perangkat_keras is not string %}
            <ul>
              {% for h in data.media_pembelajaran.perangkat_keras %}
                <li>{{ h }}</li>
              {% endfor %}
            </ul>
          {% else %}
            {{ data.media_pembelajaran.perangkat_keras or '-' }}
          {% endif %}
        {% else %}
          -
        {% endif %}
      </td>
    </tr>
  </table>

  <h2>Daftar Referensi</h2>
  <table>
    <tr>
      <th width="50%">Referensi Utama</th>
      <th width="50%">Referensi Pendukung</th>
    </tr>
    <tr>
      <td>
        {% if data.referensi is mapping %}
          {% if data.referensi.utama is iterable and data.referensi.utama is not string %}
            <ol>
              {% for u in data.referensi.utama %}
                <li>{{ u }}</li>
              {% endfor %}
            </ol>
          {% else %}
            {{ data.referensi.utama or '-' }}
          {% endif %}
        {% else %}
          {% if data.referensi is iterable and data.referensi is not string %}
            <ol>
              {% for ref in data.referensi %}
              <li>
                {% if ref is string %}
                  {{ ref }}
                {% elif ref is mapping %}
                  {{ ref.judul if ref.judul else '' }}{{ ', ' + ref.pengarang if ref.pengarang else '' }}{{ ' (' + ref.tahun|string + ')' if ref.tahun else '' }}
                {% endif %}
              </li>
              {% endfor %}
            </ol>
          {% else %}
            {{ data.referensi or '-' }}
          {% endif %}
        {% endif %}
      </td>
      <td>
        {% if data.referensi is mapping %}
          {% if data.referensi.pendukung is iterable and data.referensi.pendukung is not string %}
            <ol>
              {% for p in data.referensi.pendukung %}
                <li>{{ p }}</li>
              {% endfor %}
            </ol>
          {% else %}
            {{ data.referensi.pendukung or '-' }}
          {% endif %}
        {% else %}
          -
        {% endif %}
      </td>
    </tr>
  </table>

  <h2>Rencana Kegiatan Pembelajaran Mingguan</h2>
  <table>
    <tr>
      <th width="6%">Mg Ke-</th>
      <th width="15%">Sub-CP-MK (Kemampuan Akhir yg Diharapkan)</th>
      <th width="20%">Materi Pembelajaran</th>
      <th width="14%">Bentuk & Metode Pembelajaran</th>
      <th width="11%">Estimasi Waktu</th>
      <th width="13%">Pengalaman Belajar Mahasiswa</th>
      <th width="13%">Kriteria & Bentuk Penilaian</th>
      <th width="8%">Bobot (%)</th>
    </tr>
    {% for r in data.rencana_pembelajaran %}
    {% if r is mapping %}
    <tr>
      <td class="text-center" style="font-weight: bold;">{{ r.minggu_ke }}</td>
      <td><strong>{{ r.sub_cpmk_kode }}</strong><br/>{{ r.sub_cpmk_deskripsi or '' }}</td>
      <td>{{ r.materi }}</td>
      <td>{{ r.metode }}</td>
      <td>{{ r.estimasi_waktu or '' }}</td>
      <td>{{ r.pengalaman_belajar or '' }}</td>
      <td>{{ r.kriteria_penilaian or '' }}</td>
      <td class="text-center">{{ r.bobot or '0' }}%</td>
    </tr>
    {% endif %}
    {% endfor %}
  </table>

</body>
</html>
""")


def generate_docx(rps_data: dict, output_path: str, course_cpls: list, brand_name: str, brand_logo: str, prodi_name: str):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    identitas = rps_data.get('identitas') or {}

    # Cover Header Text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p.add_run(f"{brand_name}\nPROGRAM STUDI {prodi_name.upper()}\n\n")
    run_inst.bold = True
    run_inst.font.size = Pt(13)

    run_title = p.add_run("RENCANA PEMBELAJARAN SEMESTER (RPS)")
    run_title.bold = True
    run_title.font.size = Pt(14)

    # Identitas Table with Otorisasi block
    doc.add_heading('A. Identitas & Otorisasi', level=2)
    table = doc.add_table(rows=8, cols=4, style='Table Grid')
    
    table.rows[0].cells[0].text = 'Nama Mata Kuliah'
    table.rows[0].cells[1].text = identitas.get('nama_mata_kuliah', '')
    table.rows[0].cells[2].text = 'Kode MK'
    table.rows[0].cells[3].text = identitas.get('kode_mata_kuliah', '')
    
    table.rows[1].cells[0].text = 'Bobot (SKS)'
    table.rows[1].cells[1].text = f"{identitas.get('sks', '')} SKS"
    table.rows[1].cells[2].text = 'Semester'
    table.rows[1].cells[3].text = str(identitas.get('semester', ''))
    
    table.rows[2].cells[0].text = 'Tanggal Penyusunan'
    table.rows[2].cells[1].text = identitas.get('tanggal_penyusunan', '') or '-'
    table.rows[2].cells[2].text = 'No Dokumen'
    table.rows[2].cells[3].text = identitas.get('no_dokumen', '') or '-'
    
    # Merge row 3 for Otorisasi Section Header
    cell_oto = table.rows[3].cells[0]
    cell_oto.text = 'OTORISASI'
    for col in range(1, 4):
        cell_oto.merge(table.rows[3].cells[col])
        
    table.rows[4].cells[0].text = 'Koordinator Pengembang RPS'
    table.rows[4].cells[2].text = 'Koordinator Rumpun MK'
    table.rows[4].cells[3].text = 'Ka PRODI'
    
    table.rows[4].cells[0].merge(table.rows[4].cells[1])
    table.rows[5].cells[0].merge(table.rows[5].cells[1])
    
    table.rows[5].cells[0].text = f"Tanda tangan\n\n\n{identitas.get('koordinator_pengembang_rps', '') or '-'}"
    table.rows[5].cells[2].text = f"Tanda tangan\n\n\n{identitas.get('koordinator_rmk', '') or '-'}"
    table.rows[5].cells[3].text = f"Tanda tangan\n\n\n{identitas.get('ka_prodi', '') or '-'}"

    # Dosen Pengampu (Row 6)
    table.rows[6].cells[0].text = 'Dosen Pengampu'
    table.rows[6].cells[1].merge(table.rows[6].cells[2])
    table.rows[6].cells[1].merge(table.rows[6].cells[3])
    dosen_list = rps_data.get('dosen_pengampu') or []
    if isinstance(dosen_list, list):
        dp_text = ", ".join([str(d.get('nama', '')).strip() if isinstance(d, dict) else str(d) for d in dosen_list if d])
    else:
        dp_text = str(dosen_list)
    table.rows[6].cells[1].text = dp_text or '-'

    # Mata Kuliah Prasyarat (Row 7)
    table.rows[7].cells[0].text = 'Mata Kuliah Prasyarat'
    table.rows[7].cells[1].merge(table.rows[7].cells[2])
    table.rows[7].cells[1].merge(table.rows[7].cells[3])
    table.rows[7].cells[1].text = rps_data.get('prasyarat', '-') or '-'

    # CPL-PRODI Mapped Table
    # Capaian Pembelajaran (CP)
    doc.add_heading('B. Capaian Pembelajaran (CP)', level=2)
    
    # 1. CPL-PRODI Table
    p_cpl = doc.add_paragraph()
    p_cpl.add_run('1. Capaian Pembelajaran Lulusan Program Studi (CPL-PRODI) Yang Dibebankan Pada Mata Kuliah').bold = True
    cpl_table = doc.add_table(rows=1, cols=2, style='Table Grid')
    cpl_table.rows[0].cells[0].text = 'Kode CPL'
    cpl_table.rows[0].cells[1].text = 'Deskripsi CPL yang Dibebankan'
    
    for cpl in course_cpls:
        row = cpl_table.add_row()
        row.cells[0].text = cpl.get('kode', '')
        row.cells[1].text = cpl.get('deskripsi', '')

    doc.add_paragraph() # Spacer

    # 2. CPMK Table
    p_cpmk = doc.add_paragraph()
    p_cpmk.add_run('2. Capaian Pembelajaran Mata Kuliah (CPMK)').bold = True
    cpmk_table = doc.add_table(rows=1, cols=3, style='Table Grid')
    cpmk_table.rows[0].cells[0].text = 'Kode CPMK'
    cpmk_table.rows[0].cells[1].text = 'Deskripsi CPMK'
    cpmk_table.rows[0].cells[2].text = 'CPL Terkait'
    
    for c in (rps_data.get('cpmk') or []):
        if not isinstance(c, dict):
            continue
        row = cpmk_table.add_row()
        row.cells[0].text = c.get('kode', '')
        row.cells[1].text = c.get('deskripsi', '')
        cpl_prodi = c.get('cpl_prodi', [])
        row.cells[2].text = ', '.join(cpl_prodi) if isinstance(cpl_prodi, list) else str(cpl_prodi or '')

    doc.add_paragraph() # Spacer

    # Deskripsi
    doc.add_heading('C. Deskripsi Singkat MK', level=2)
    doc.add_paragraph(rps_data.get('deskripsi_mata_kuliah', ''))

    # Bahan Kajian
    doc.add_heading('E. Bahan Kajian / Pokok Bahasan', level=2)
    bahan_kajian = rps_data.get('bahan_kajian', [])
    if isinstance(bahan_kajian, list):
        for bk in bahan_kajian:
            doc.add_paragraph(str(bk), style='List Bullet')
    else:
        doc.add_paragraph(str(bahan_kajian or '-'))

    # Media Pembelajaran Table
    doc.add_heading('F. Media Pembelajaran', level=2)
    media_table = doc.add_table(rows=2, cols=2, style='Table Grid')
    media_table.rows[0].cells[0].text = 'Perangkat Lunak (Software)'
    media_table.rows[0].cells[1].text = 'Perangkat Keras (Hardware)'
    
    media_data = rps_data.get('media_pembelajaran') or {}
    if isinstance(media_data, dict):
        soft = '\n'.join(media_data.get('perangkat_lunak', [])) if isinstance(media_data.get('perangkat_lunak'), list) else str(media_data.get('perangkat_lunak', ''))
        hard = '\n'.join(media_data.get('perangkat_keras', [])) if isinstance(media_data.get('perangkat_keras'), list) else str(media_data.get('perangkat_keras', ''))
    else:
        soft = '\n'.join(media_data) if isinstance(media_data, list) else str(media_data)
        hard = '-'
    media_table.rows[1].cells[0].text = soft
    media_table.rows[1].cells[1].text = hard

    # Referensi Table
    doc.add_heading('G. Daftar Referensi', level=2)
    ref_table = doc.add_table(rows=2, cols=2, style='Table Grid')
    ref_table.rows[0].cells[0].text = 'Referensi Utama'
    ref_table.rows[0].cells[1].text = 'Referensi Pendukung'
    
    ref_data = rps_data.get('referensi') or {}
    if isinstance(ref_data, dict):
        utama = '\n'.join(ref_data.get('utama', [])) if isinstance(ref_data.get('utama'), list) else str(ref_data.get('utama', ''))
        pendukung = '\n'.join(ref_data.get('pendukung', [])) if isinstance(ref_data.get('pendukung'), list) else str(ref_data.get('pendukung', ''))
    else:
        utama = '\n'.join([f"{r.get('pengarang','')}. {r.get('tahun','')}. {r.get('judul','')}" for r in ref_data if isinstance(r, dict)])
        pendukung = '-'
    ref_table.rows[1].cells[0].text = utama
    ref_table.rows[1].cells[1].text = pendukung

    # Rencana Pembelajaran Table Columns
    doc.add_heading('H. Rencana Kegiatan Pembelajaran Mingguan', level=2)
    rp_table = doc.add_table(rows=1, cols=8, style='Table Grid')
    
    headers = [
        'Mg Ke-',
        'Sub-CP-MK (Kemampuan Akhir yg Diharapkan)',
        'Materi Pembelajaran',
        'Bentuk & Metode Pembelajaran',
        'Estimasi Waktu',
        'Pengalaman Belajar Mahasiswa',
        'Kriteria & Bentuk Penilaian',
        'Bobot (%)'
    ]
    for j, h in enumerate(headers):
        rp_table.rows[0].cells[j].text = h
        
    for rp in (rps_data.get('rencana_pembelajaran') or []):
        if not isinstance(rp, dict):
            continue
        row = rp_table.add_row()
        row.cells[0].text = str(rp.get('minggu_ke', '')) or ''
        
        sub_cpmk_text = f"{rp.get('sub_cpmk_kode', '')}\n{rp.get('sub_cpmk_deskripsi', '')}"
        row.cells[1].text = sub_cpmk_text.strip() or ''
        row.cells[2].text = rp.get('materi', '') or ''
        row.cells[3].text = str(rp.get('metode', '')) or ''
        row.cells[4].text = rp.get('estimasi_waktu', '') or ''
        row.cells[5].text = rp.get('pengalaman_belajar', '') or ''
        row.cells[6].text = rp.get('kriteria_penilaian', '') or ''
        row.cells[7].text = f"{rp.get('bobot', '0')}%"

    doc.save(output_path)


@router.post("/{rps_id}")
@router.get("/{rps_id}")
async def export_rps(
    rps_id: int,
    export_format: str = "pdf",
    db: AsyncSession = Depends(get_db),
):
    from app.models import MataKuliah, Prodi
    result = await db.execute(select(RPS).where(RPS.id == rps_id))
    rps = result.scalar_one_or_none()
    if not rps:
        raise HTTPException(status_code=404, detail="RPS not found")
        
    # Ambil deskripsi mata kuliah dari database secara dinamis agar tidak hardcoded
    mk_result = await db.execute(select(MataKuliah).where(MataKuliah.id == rps.mata_kuliah_id))
    mk = mk_result.scalar_one_or_none()
    deskripsi_mk = mk.deskripsi if (mk and mk.deskripsi) else ""
    
    # Ambil data CPL Prodi & subset CPL yang dipetakan pada mata kuliah ini
    prodi_result = await db.execute(select(Prodi).where(Prodi.id == rps.prodi_id))
    prodi = prodi_result.scalar_one_or_none()
    
    all_cpls = prodi.capaian_pembelajaran_lulusan if (prodi and prodi.capaian_pembelajaran_lulusan) else []
    course_cpl_codes = mk.cpl_prodi or [] if mk else []
    course_cpls = [cpl for cpl in all_cpls if cpl.get("kode") in course_cpl_codes]
    if not course_cpls:
        course_cpls = all_cpls
        
    # Retrieve prerequisites (prasyarat)
    prasyarat_text = "-"
    if mk and mk.prasyarat:
        if isinstance(mk.prasyarat, list):
            prasyarat_text = ", ".join(mk.prasyarat)
        else:
            prasyarat_text = str(mk.prasyarat)

    prodi_name = prodi.nama if prodi else ""
    
    rps_data = {
        "identitas": rps.identitas or {},
        "deskripsi_mata_kuliah": deskripsi_mk or "",
        "cpmk": rps.cpmk or [],
        "sub_cpmk": rps.sub_cpmk or [],
        "rencana_pembelajaran": rps.rencana_pembelajaran or [],
        "metode_pembelajaran": rps.metode_pembelajaran or [],
        "media_pembelajaran": rps.media_pembelajaran or [],
        "penilaian": rps.penilaian or [],
        "referensi": rps.referensi or [],
        "dosen_pengampu": rps.dosen_pengampu or [],
        "prasyarat": prasyarat_text or "-",
    }
    
    os.makedirs(settings.EXPORT_DIR, exist_ok=True)
    filename = f"RPS_{rps.kode}_{rps.semester}_{rps.tahun_akademik}"
    filename = filename.replace("/", "-").replace("\\", "-")
    
    if export_format == "docx":
        filepath = os.path.join(settings.EXPORT_DIR, f"{filename}.docx")
        generate_docx(rps_data, filepath, course_cpls, brand_name, brand_logo, prodi_name)
        return FileResponse(filepath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"{filename}.docx")
    
    elif export_format == "pdf":
        html_content = RPS_HTML_TEMPLATE.render(
            data=rps_data,
            course_cpls=course_cpls,
            brand_name=brand_name,
            brand_logo=brand_logo,
            prodi_name=prodi_name
        )
        
        try:
            from xhtml2pdf import pisa
            pdf_filename = f"{filename}.pdf"
            filepath = os.path.join(settings.EXPORT_DIR, pdf_filename)
            with open(filepath, "w+b") as result_file:
                pisa_status = pisa.CreatePDF(html_content, dest=result_file)
                if pisa_status.err:
                    raise Exception("Gagal mengonversi HTML ke PDF")
            
            return FileResponse(
                filepath,
                media_type="application/pdf",
                filename=pdf_filename,
                headers={"Content-Disposition": f"attachment; filename={pdf_filename}"}
            )
        except ImportError:
            # Fallback: jika library xhtml2pdf belum terinstall di server, return file html
            html_filename = f"{filename}.html"
            filepath = os.path.join(settings.EXPORT_DIR, html_filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(html_content)
            return FileResponse(
                filepath,
                media_type="text/html",
                filename=html_filename,
                headers={"Content-Disposition": f"attachment; filename={html_filename}"}
            )
    
    else:
        raise HTTPException(status_code=400, detail=f"Format {export_format} tidak didukung")